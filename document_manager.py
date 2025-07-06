"""文档管理模块

该模块整合了文档加载、处理和向量化功能，包括：
- 文档加载和解析（支持多种格式）
- 文档预处理和分块
- 向量化（embedding）
- FAISS索引构建和管理
- 文档-块映射关系维护
- 缓存管理

重构优化：
- 模块化设计，减少冗余代码
- 所有文档解析后输出markdown格式
- 文档分块后输出json格式
- 专注于文档处理，不包含检索功能
"""

import os
import re
import json
import pickle
import hashlib
import asyncio
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any, Union, Callable
from functools import lru_cache
from abc import ABC, abstractmethod

import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import PyPDF2
from docx import Document
import pandas as pd
import openpyxl
import win32com.client
import pythoncom

from config import config


class DocumentFormat(ABC):
    """文档格式抽象基类，定义了文档格式处理的接口"""
    
    @abstractmethod
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载文档并返回内容"""
        pass
    
    @abstractmethod
    def to_markdown(self, content: Any) -> str:
        """将文档内容转换为Markdown格式"""
        pass


class PDFFormat(DocumentFormat):
    """PDF文档格式处理类"""
    
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载PDF文件
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            包含文档内容的字典
        """
        content = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text:
                        content.append({
                            'page': i + 1,
                            'text': text
                        })
        except Exception as e:
            raise IOError(f"加载PDF文件失败: {str(e)}")
                        
        return {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'content': content,
            'total_pages': len(content)
        }
    
    def to_markdown(self, content: List[Dict[str, Any]]) -> str:
        """将PDF内容转换为Markdown格式
        
        Args:
            content: PDF内容列表，每项包含页码和文本
            
        Returns:
            Markdown格式的文本
        """
        markdown = []
        for page in content:
            markdown.append(f"## 第{page['page']}页\n\n{page['text']}\n")
        return "\n".join(markdown)


class DocxFormat(DocumentFormat):
    """DOCX文档格式处理类"""
    
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            包含文档内容的字典
        """
        doc = Document(file_path)
        paragraphs = []
        tables = []
        
        # 提取段落
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append({
                    'index': i,
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })
        
        # 提取表格
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': i,
                'data': table_data
            })
                
        return {
            'file_path': str(file_path),
            'file_type': 'docx',
            'paragraphs': paragraphs,
            'tables': tables
        }
    
    def to_markdown(self, content: Dict[str, Any]) -> str:
        """将Word文档内容转换为Markdown格式
        
        Args:
            content: Word文档内容字典
            
        Returns:
            Markdown格式的文本
        """
        markdown = []
        
        # 处理段落
        for para in content['paragraphs']:
            style = para['style']
            text = para['text']
            
            # 根据样式转换为对应的Markdown格式
            if style.startswith('Heading'):
                level = int(style[-1]) if style[-1].isdigit() else 1
                markdown.append(f"{'#' * level} {text}\n")
            else:
                markdown.append(f"{text}\n\n")
        
        # 处理表格
        for table in content['tables']:
            if not table['data']:
                continue
                
            # 创建表头
            header = table['data'][0]
            markdown.append('|' + '|'.join(header) + '|')
            markdown.append('|' + '|'.join(['---'] * len(header)) + '|')
            
            # 添加表格内容
            for row in table['data'][1:]:
                markdown.append('|' + '|'.join(row) + '|')
            
            markdown.append('\n')
            
        return "\n".join(markdown)


class DocFormat(DocumentFormat):
    """DOC文档格式处理类"""
    
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载.doc格式的Word文档
        
        Args:
            file_path: DOC文档路径
            
        Returns:
            包含文档内容的字典
        """
        try:
            # 初始化COM
            pythoncom.CoInitialize()
            
            # 创建Word应用程序实例
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                # 打开文档
                doc = word.Documents.Open(str(file_path.absolute()))
                # 获取文本内容
                content = doc.Content.Text
                # 关闭文档
                doc.Close()
            finally:
                # 退出Word应用程序
                word.Quit()
                # 释放COM
                pythoncom.CoUninitialize()
            
            return {
                'file_path': str(file_path),
                'file_type': 'doc',
                'content': content
            }
        except Exception as e:
            raise IOError(f"加载.doc文件失败: {str(e)}")
    
    def to_markdown(self, content: str) -> str:
        """将DOC文档内容转换为Markdown格式
        
        Args:
            content: DOC文档内容
            
        Returns:
            Markdown格式的文本
        """
        # 简单处理，将段落分隔并添加空行
        paragraphs = content.split('\r\n')
        return '\n\n'.join(paragraphs)


class ExcelFormat(DocumentFormat):
    """Excel文档格式处理类"""
    
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载Excel文件
        
        Args:
            file_path: Excel文件路径
            
        Returns:
            包含文档内容的字典
        """
        sheet_contents = {}  # 存储每个工作表的内容
        
        try:
            # 使用pandas读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheet_contents[sheet_name] = {
                    'dataframe': df,
                    'raw_text': df.to_string()
                }
        except Exception as e:
            # 如果pandas失败，使用openpyxl作为备选
            try:
                wb = openpyxl.load_workbook(file_path, data_only=True)
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.rows:
                        row_data = [str(cell.value) if cell.value is not None else '' for cell in row]
                        rows.append(row_data)
                    
                    # 创建DataFrame
                    if rows:
                        df = pd.DataFrame(rows[1:], columns=rows[0] if rows else None)
                        sheet_contents[sheet_name] = {
                            'dataframe': df,
                            'raw_text': '\n'.join(['\t'.join(row) for row in rows])
                        }
            except Exception as nested_e:
                raise IOError(f"加载Excel文件失败: {str(e)}，备选方法也失败: {str(nested_e)}")
                
        return {
            'file_path': str(file_path),
            'file_type': 'excel',
            'sheets': sheet_contents
        }
    
    def to_markdown(self, content: Dict[str, Any]) -> str:
        """将Excel内容转换为Markdown格式
        
        Args:
            content: Excel内容字典
            
        Returns:
            Markdown格式的文本
        """
        markdown = []
        
        for sheet_name, sheet_data in content['sheets'].items():
            markdown.append(f"## 工作表: {sheet_name}\n")
            
            df = sheet_data['dataframe']
            if not df.empty:
                # 创建表头
                headers = df.columns.tolist()
                markdown.append('|' + '|'.join(str(h) for h in headers) + '|')
                markdown.append('|' + '|'.join(['---'] * len(headers)) + '|')
                
                # 添加表格内容（限制行数以避免过大）
                max_rows = min(100, len(df))
                for _, row in df.head(max_rows).iterrows():
                    markdown.append('|' + '|'.join(str(cell) for cell in row) + '|')
                
                if len(df) > max_rows:
                    markdown.append(f"\n*注: 表格显示了 {max_rows} 行，共 {len(df)} 行*\n")
            else:
                markdown.append("*空工作表*\n")
            
            markdown.append("\n")
            
        return "\n".join(markdown)


class TextFormat(DocumentFormat):
    """文本文档格式处理类"""
    
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载文本文件
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            包含文档内容的字典
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            return {
                'file_path': str(file_path),
                'file_type': 'txt',
                'content': content
            }
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    return {
                        'file_path': str(file_path),
                        'file_type': 'txt',
                        'content': content,
                        'encoding': encoding
                    }
                except UnicodeDecodeError:
                    continue
            raise IOError(f"无法解码文本文件: {file_path}")
    
    def to_markdown(self, content: str) -> str:
        """将文本内容转换为Markdown格式
        
        Args:
            content: 文本内容
            
        Returns:
            Markdown格式的文本
        """
        # 文本文件直接返回，不需要特殊处理
        return content


class CSVFormat(DocumentFormat):
    """CSV文档格式处理类"""
    
    def load(self, file_path: Path) -> Dict[str, Any]:
        """加载CSV文件
        
        Args:
            file_path: CSV文件路径
            
        Returns:
            包含文档内容的字典
        """
        try:
            df = pd.read_csv(file_path)
            return {
                'file_path': str(file_path),
                'file_type': 'csv',
                'dataframe': df,
                'raw_text': df.to_string()
            }
        except Exception as e:
            # 尝试不同的编码和分隔符
            for encoding in ['utf-8', 'gbk', 'latin-1']:
                for sep in [',', ';', '\t']:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                        return {
                            'file_path': str(file_path),
                            'file_type': 'csv',
                            'dataframe': df,
                            'raw_text': df.to_string(),
                            'encoding': encoding,
                            'separator': sep
                        }
                    except Exception:
                        continue
            raise IOError(f"加载CSV文件失败: {str(e)}")
    
    def to_markdown(self, content: Dict[str, Any]) -> str:
        """将CSV内容转换为Markdown格式
        
        Args:
            content: CSV内容字典
            
        Returns:
            Markdown格式的文本
        """
        markdown = []
        df = content['dataframe']
        
        if not df.empty:
            # 创建表头
            headers = df.columns.tolist()
            markdown.append('|' + '|'.join(str(h) for h in headers) + '|')
            markdown.append('|' + '|'.join(['---'] * len(headers)) + '|')
            
            # 添加表格内容（限制行数以避免过大）
            max_rows = min(100, len(df))
            for _, row in df.head(max_rows).iterrows():
                markdown.append('|' + '|'.join(str(cell) for cell in row) + '|')
            
            if len(df) > max_rows:
                markdown.append(f"\n*注: 表格显示了 {max_rows} 行，共 {len(df)} 行*")
        else:
            markdown.append("*空CSV文件*")
        
        return "\n".join(markdown)


class DocumentLoader:
    """文档加载器，支持多种文件格式的加载和转换为Markdown"""
    
    def __init__(self):
        """初始化文档加载器"""
        self.format_handlers = {
            '.pdf': PDFFormat(),
            '.docx': DocxFormat(),
            '.doc': DocFormat(),
            '.xlsx': ExcelFormat(),
            '.xls': ExcelFormat(),
            '.txt': TextFormat(),
            '.csv': CSVFormat()
        }
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """加载单个文档并转换为Markdown格式
        
        Args:
            file_path: 文件路径
            
        Returns:
            包含文档内容和Markdown格式的字典
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        extension = file_path.suffix.lower()
        if extension not in self.format_handlers:
            raise ValueError(f"不支持的文件格式: {extension}")
            
        handler = self.format_handlers[extension]
        doc_info = handler.load(file_path)
        
        # 转换为Markdown格式
        if extension == '.pdf':
            markdown = handler.to_markdown(doc_info['content'])
        elif extension in ['.docx', '.xlsx', '.xls', '.csv']:
            markdown = handler.to_markdown(doc_info)
        else:  # .doc, .txt
            markdown = handler.to_markdown(doc_info['content'])
        
        # 添加Markdown格式到结果中
        doc_info['markdown'] = markdown
        
        return doc_info
    
    def load_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """加载目录中的所有支持格式的文档
        
        Args:
            directory_path: 目录路径
            
        Returns:
            文档内容列表
        """
        directory_path = Path(directory_path)
        if not directory_path.exists():
            raise FileNotFoundError(f"目录不存在: {directory_path}")
            
        documents = []
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.format_handlers:
                try:
                    doc = self.load_document(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"加载文件 {file_path} 时出错: {str(e)}")
                    
        return documents


class VectorCache:
    """向量缓存管理类"""
    
    def __init__(self, cache_file: str):
        """初始化向量缓存
        
        Args:
            cache_file: 缓存文件路径
        """
        self.cache_file = cache_file
        self.cache: Dict[str, np.ndarray] = {}
        self.load_cache()
    
    def load_cache(self):
        """加载缓存"""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'rb') as f:
                    self.cache = pickle.load(f)
            except Exception:
                self.cache = {}
    
    def save_cache(self):
        """保存缓存"""
        try:
            # 确保缓存目录存在
            os.makedirs(os.path.dirname(self.cache_file), exist_ok=True)
            with open(self.cache_file, 'wb') as f:
                pickle.dump(self.cache, f)
        except Exception as e:
            print(f"保存缓存失败: {str(e)}")
    
    def get(self, key: str) -> Optional[np.ndarray]:
        """获取缓存的向量
        
        Args:
            key: 缓存键
            
        Returns:
            缓存的向量或None
        """
        return self.cache.get(key)
    
    def set(self, key: str, vector: np.ndarray):
        """设置向量缓存
        
        Args:
            key: 缓存键
            vector: 向量数据
        """
        self.cache[key] = vector


class DocumentChunker:
    """文档分块器，将文档切分成较小的块"""
    
    def __init__(self, config_obj=None):
        """初始化文档分块器
        
        Args:
            config_obj: 配置对象，默认使用全局配置
        """
        self.config = config_obj or config
    
    def chunk_document(self, text: str, max_chars: int = None, overlap: int = None, 
                      is_excel: bool = False, file_name: str = None) -> List[str]:
        """将长文档切分成较小的块，使用滑动窗口确保上下文连贯性
        
        Args:
            text: 要切分的文本
            max_chars: 每个块的最大字符数
            overlap: 相邻块之间的重叠字符数
            is_excel: 是否为Excel文件
            file_name: 文件名
        
        Returns:
            切分后的文本块列表
        """
        # 使用配置中的默认值
        if max_chars is None:
            max_chars = self.config.system_config["max_chunk_chars"]
        if overlap is None:
            overlap = self.config.system_config["chunk_overlap"]
        
        # 根据文件类型设置分块参数
        if is_excel:
            max_chars = self.config.system_config["excel_max_chunk_chars"]
            overlap = self.config.system_config["excel_chunk_overlap"]
        
        # 文件名前缀
        file_prefix = ""
        if file_name:
            file_prefix = f"[文件：{file_name}]\n"
        
        if len(text) <= max_chars:
            # 如果整个文本小于一个块的大小，直接添加文件名前缀返回
            if file_name and not text.startswith(file_prefix):
                return [f"{file_prefix}{text}"]
            return [text]
        
        # 移除可能存在的文件名前缀，以避免重复
        if file_name and text.startswith(file_prefix):
            text = text[len(file_prefix):]
        
        chunks = []
        start = 0
        last_end = 0

        while start < len(text):
            end = min(start + max_chars, len(text))
            
            if end < len(text):
                sentence_ends = [m.end() for m in re.finditer(r'[。！？.!?]\s*', text[start:end])]
                
                if sentence_ends:
                    end = start + sentence_ends[-1]
                else:  # 如果没有找到，尝试在单词或标点处切分
                    last_space = text[start:end].rfind(' ')
                    last_punct = max(text[start:end].rfind('，'), text[start:end].rfind(','))
                    cut_point = max(last_space, last_punct)
                    
                    if cut_point > 0:
                        end = start + cut_point + 1
            
            # 每个chunk都添加文件名前缀
            chunk = text[start:end]
            if file_name:
                chunk = f"{file_prefix}{chunk}"
            
            chunks.append(chunk)
            
            if end <= last_end:
                end = min(last_end + 1, len(text))
                chunk = text[start:end]
                if file_name:
                    chunk = f"{file_prefix}{chunk}"
                chunks[-1] = chunk
                
                if end >= len(text):
                    break
            
            last_end = end
            start = end - overlap
            
            if start < 0:
                start = 0
                
            if start >= end:
                start = end
                
            if start >= len(text):
                break
        
        return chunks
    
    def save_chunks_to_json(self, chunks: List[str], file_name: str, source_path: str, markdown_content: str = None):
        """保存文档块到JSON文件
        
        Args:
            chunks: 文档块列表
            file_name: 文件名
            source_path: 源文件路径
            markdown_content: Markdown格式的文档内容
        """
        chunks_dir = Path(self.config.system_config["chunks_dir"])
        chunks_dir.mkdir(exist_ok=True)
        
        chunks_data = {
            "file_name": file_name,
            "file_path": source_path,
            "total_chunks": len(chunks),
            "chunks": [
                {
                    "chunk_id": j,
                    "content": chunk,
                    "length": len(chunk)
                }
                for j, chunk in enumerate(chunks)
            ],
            "created_at": str(pd.Timestamp.now())
        }
        
        # 如果提供了Markdown内容，添加到JSON中
        if markdown_content:
            chunks_data["markdown"] = markdown_content
        
        # 将文件名中的特殊字符替换为下划线，作为JSON文件名
        safe_filename = re.sub(r'[^\w\-_.]', '_', file_name)
        json_path = chunks_dir / f"{safe_filename}.json"
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(chunks_data, f, ensure_ascii=False, indent=2)
        
        return json_path


class EmbeddingModel:
    """嵌入模型类，负责文本向量化"""
    
    def __init__(self, config_obj=None):
        """初始化嵌入模型
        
        Args:
            config_obj: 配置对象，默认使用全局配置
        """
        self.config = config_obj or config
        self.model: Optional[SentenceTransformer] = None
        self.embedding_cache = VectorCache(
            str(Path(self.config.system_config["cache_dir"]) / self.config.system_config["embedding_cache_file"])
        )
    
    def load_model(self) -> SentenceTransformer:
        """加载或初始化模型
        
        Returns:
            加载的模型实例
        """
        local_model_path = self.config.system_config["local_model_path"]
        
        if os.path.exists(local_model_path):
            print(f"从本地加载模型: {local_model_path}")
            model = SentenceTransformer(local_model_path)
        else:
            print(f"本地模型不存在，从网络加载: moka-ai/m3e-base")
            model = SentenceTransformer('moka-ai/m3e-base')
            print(f"保存模型到本地: {local_model_path}")
            model.save(local_model_path)
        
        # 尝试使用GPU
        try:
            model = model.to('cuda')
            print('模型已转移到GPU')
        except Exception:
            print('未检测到可用GPU，使用CPU')
        
        self.model = model
        return model
    
    @lru_cache(maxsize=1000)
    def get_cached_embeddings(self, text: str) -> np.ndarray:
        """缓存文本的嵌入向量
        
        Args:
            text: 输入文本
            
        Returns:
            嵌入向量
        """
        # 首先检查持久化缓存
        cache_key = hashlib.md5(text.encode()).hexdigest()
        cached_vector = self.embedding_cache.get(cache_key)
        if cached_vector is not None:
            return cached_vector

        # 如果缓存中没有，计算新的嵌入向量
        if self.model is None:
            raise RuntimeError("模型未初始化")
        
        vector = self.model.encode(text, convert_to_tensor=False)
        
        # 保存到持久化缓存
        self.embedding_cache.set(cache_key, vector)
        self.embedding_cache.save_cache()
        return vector
    
    def get_embeddings(self, texts: List[str]) -> np.ndarray:
        """获取文本嵌入向量
        
        Args:
            texts: 文本列表
        
        Returns:
            numpy数组形式的嵌入向量
        """
        if self.model is None:
            raise RuntimeError("模型未初始化")
        
        embeddings = self.model.encode(texts, normalize_embeddings=True)
        return np.array(embeddings)


class FaissIndexManager:
    """FAISS索引管理类"""
    
    def __init__(self, config_obj=None):
        """初始化FAISS索引管理器
        
        Args:
            config_obj: 配置对象，默认使用全局配置
        """
        self.config = config_obj or config
        self.index: Optional[faiss.Index] = None
    
    def create_index(self, embeddings: np.ndarray) -> faiss.Index:
        """创建FAISS索引（兼容模式：优先使用GPU，失败时回退到CPU）
        
        Args:
            embeddings: 嵌入向量数组
            
        Returns:
            FAISS索引
        """
        dimension = embeddings.shape[1]
        nlist = min(100, len(embeddings) // 10)  # 聚类中心数量
        
        if len(embeddings) > nlist and nlist > 0:
            # 使用IVF索引类型，提高搜索效率
            quantizer = faiss.IndexFlatL2(dimension)
            index = faiss.IndexIVFFlat(quantizer, dimension, nlist)
            
            # 训练索引
            if not index.is_trained:
                print("训练索引...")
                index.train(embeddings)
        else:
            # 对于小数据集，使用简单的平面索引
            index = faiss.IndexFlatL2(dimension)
        
        # 添加向量到索引（在GPU转移之前）
        index.add(embeddings)
        
        # 兼容模式：尝试将索引转移到GPU，失败时继续使用CPU
        gpu_index = self._try_move_to_gpu(index)
        if gpu_index is not None:
            self.index = gpu_index
            print("FAISS索引已成功转移到GPU")
        else:
            self.index = index
            print("FAISS索引使用CPU模式")
        
        return self.index
    
    def _try_move_to_gpu(self, index: faiss.Index) -> Optional[faiss.Index]:
        """尝试将FAISS索引转移到GPU（兼容模式）
        
        Args:
            index: CPU上的FAISS索引
            
        Returns:
            GPU索引（成功时）或None（失败时）
        """
        try:
            import torch
            # 检查CUDA是否可用
            if not torch.cuda.is_available():
                return None
            
            # 检查是否有faiss-gpu支持
            if not hasattr(faiss, 'StandardGpuResources'):
                print("警告：检测到faiss-cpu版本，无法使用GPU加速")
                return None
            
            # 尝试创建GPU资源
            gpu_res = faiss.StandardGpuResources()
            
            # 尝试转移到GPU
            gpu_index = faiss.index_cpu_to_gpu(gpu_res, 0, index)
            
            # 验证GPU索引是否工作正常
            if gpu_index.ntotal != index.ntotal:
                print("警告：GPU索引验证失败，回退到CPU")
                return None
                
            return gpu_index
            
        except ImportError:
            print("PyTorch未安装或CUDA不可用")
            return None
        except AttributeError as e:
            print(f"FAISS GPU功能不可用: {str(e)}")
            return None
        except Exception as e:
            print(f"GPU转移失败，回退到CPU: {str(e)}")
            return None
    
    def save_index(self, index_path: str):
        """保存索引到文件
        
        Args:
            index_path: 索引文件路径
        """
        if self.index is not None:
            print(f"保存索引到本地: {index_path}")
            faiss.write_index(self.index, index_path)
    
    def load_index(self, index_path: str) -> Optional[faiss.Index]:
        """加载索引文件
        
        Args:
            index_path: 索引文件路径
            
        Returns:
            加载的索引或None
        """
        if not os.path.exists(index_path):
            print(f"索引文件不存在: {index_path}")
            return None
        
        try:
            # 加载索引到CPU
            cpu_index = faiss.read_index(index_path)
            
            # 兼容模式：尝试将索引转移到GPU，失败时使用CPU
            gpu_index = self._try_move_to_gpu(cpu_index)
            if gpu_index is not None:
                self.index = gpu_index
                print(f"已加载的FAISS索引已成功转移到GPU")
            else:
                self.index = cpu_index
                print(f"已加载的FAISS索引使用CPU模式")
            
            return self.index
            
        except Exception as e:
            print(f"加载索引失败: {str(e)}")
            return None


class DocumentProcessor:
    """文档处理器类，整合文档加载、处理和向量化功能"""
    
    def __init__(self, config_obj=None):
        """初始化文档处理器
        
        Args:
            config_obj: 配置对象，默认使用全局配置
        """
        self.config = config_obj or config
        self.all_chunks: List[str] = []
        self.doc_sources: List[str] = []
        self.chunks_to_document: Dict[int, int] = {}
        self.document_to_chunks: Dict[int, List[str]] = {}
        
        # 初始化组件
        self._init_components()
        
        # 初始化状态
        self.initialized = False
    
    def _init_components(self):
        """初始化各个组件"""
        # 创建缓存目录
        cache_dir = Path(self.config.system_config["cache_dir"])
        cache_dir.mkdir(exist_ok=True)
        
        # 初始化文档加载器
        self.document_loader = DocumentLoader()
        
        # 初始化文档分块器
        self.document_chunker = DocumentChunker(self.config)
        
        # 初始化嵌入模型
        self.embedding_model = EmbeddingModel(self.config)
        
        # 初始化索引管理器
        self.index_manager = FaissIndexManager(self.config)
    
    def initialize(self) -> bool:
        """初始化文档处理器（仅加载模型，不自动处理文档）
        
        Returns:
            是否成功初始化
        """
        # 加载嵌入模型
        self.embedding_model.load_model()
        
        # 初始化空的索引结构
        self.all_chunks = []
        self.doc_sources = []
        self.chunks_to_document = {}
        self.document_to_chunks = {}
        
        print("文档处理器初始化完成，等待用户上传文档")
        self.initialized = True
        return True
    
    def save_index_and_mapping(self):
        """保存索引和映射关系"""
        index_path = self.config.system_config["index_file"]
        chunks_map_path = self.config.system_config["chunks_mapping_file"]
        
        # 保存索引
        self.index_manager.save_index(index_path)
        
        # 保存映射关系和文档源
        np.save(chunks_map_path, {
            'document_to_chunks': self.document_to_chunks,
            'chunks_to_document': self.chunks_to_document,
            'doc_sources': self.doc_sources
        })
        
        # 保存嵌入缓存
        self.embedding_model.embedding_cache.save_cache()
    
    def load_index_and_mapping(self) -> bool:
        """加载索引和映射关系
        
        Returns:
            是否成功加载
        """
        index_path = self.config.system_config["index_file"]
        chunks_map_path = self.config.system_config["chunks_mapping_file"]
        
        if not (os.path.exists(index_path) and os.path.exists(chunks_map_path)):
            print(f"索引文件或映射文件不存在: {index_path}, {chunks_map_path}")
            return False
        
        try:
            # 加载映射文件
            mapping_data = np.load(chunks_map_path, allow_pickle=True).item()
            self.chunks_to_document = mapping_data.get('chunks_to_document', {})
            self.document_to_chunks = mapping_data.get('document_to_chunks', {})
            self.doc_sources = mapping_data.get('doc_sources', [])
            
            # 加载索引
            if self.index_manager.load_index(index_path):
                print(f"已加载索引和映射，包含 {len(self.doc_sources)} 个文档源")
                return True
            return False
            
        except Exception as e:
            print(f"加载索引和映射失败: {str(e)}")
            # 删除损坏的文件
            if os.path.exists(chunks_map_path):
                os.remove(chunks_map_path)
            if os.path.exists(index_path):
                os.remove(index_path)
            return False
    
    def process_documents(self) -> bool:
        """处理所有文档
        
        Returns:
            是否成功处理
        """
        print("开始处理文档...")
        
        # 加载文档
        documents, doc_sources = self._load_documents()
        
        if not documents:
            print("错误: 未能加载任何文档，请确保文档目录存在且包含有效文件")
            return False
        
        self.doc_sources = doc_sources
        
        # 处理文档分块
        self.document_to_chunks = {}
        self.chunks_to_document = {}
        self.all_chunks = []
        
        print("开始处理文档分块")
        for i, (doc, source) in enumerate(zip(documents, doc_sources)):
            # 获取文件名
            file_name = os.path.basename(source)
            # 检查是否为Excel文件
            is_excel = source.lower().endswith(('.xlsx', '.xls'))
            
            # 使用文档的Markdown格式进行分块
            markdown_content = doc.get('markdown', doc.get('content', ''))
            chunks = self.document_chunker.chunk_document(markdown_content, is_excel=is_excel, file_name=file_name)
            self.document_to_chunks[i] = chunks
            
            # 保存文档块到JSON文件
            self.document_chunker.save_chunks_to_json(chunks, file_name, source, markdown_content)
            
            for chunk in chunks:
                self.chunks_to_document[len(self.all_chunks)] = i
                self.all_chunks.append(chunk)
        
        if not self.all_chunks:
            print("错误: 没有生成任何文档块")
            return False
            
        print(f"成功生成 {len(self.all_chunks)} 个文档块")
        
        # 创建索引
        print("创建索引")
        embeddings = self.embedding_model.get_embeddings(self.all_chunks)
        self.index_manager.create_index(embeddings)
        
        # 保存索引和映射
        self.save_index_and_mapping()
        
        return True
    
    def _load_documents(self) -> Tuple[List[Dict[str, Any]], List[str]]:
        """加载uploads目录下的所有文档
        
        Returns:
            tuple: (documents, doc_sources)
                - documents: 文档内容列表
                - doc_sources: 文档源文件路径列表
        """
        # 获取uploads目录路径
        uploads_dir = Path("uploads")
        if not uploads_dir.exists():
            uploads_dir.mkdir(exist_ok=True)
            print(f"创建uploads目录: {uploads_dir}")
            return [], []
        
        # 加载所有文档
        try:
            docs = self.document_loader.load_directory(str(uploads_dir))
            doc_sources = [doc["file_path"] for doc in docs]
            return docs, doc_sources
        except Exception as e:
            print(f"加载文档时出错: {str(e)}")
            return [], []
    
    async def add_document_async(self, doc_path: str) -> bool:
        """异步增量添加文档到索引
        
        Args:
            doc_path: 文档路径
            
        Returns:
            是否成功添加
        """
        print(f"开始异步处理新上传的文档: {doc_path}")
        
        try:
            # 加载文档内容
            doc_info = self.document_loader.load_document(doc_path)
            file_name = os.path.basename(doc_path)
            is_excel = file_name.lower().endswith(('.xlsx', '.xls'))
            
            # 使用文档的Markdown格式进行分块
            markdown_content = doc_info.get('markdown', '')
            chunks = self.document_chunker.chunk_document(markdown_content, is_excel=is_excel, file_name=file_name)
            
            if not chunks:
                print("未生成任何文档块，跳过")
                return False
            
            print(f"为文档 {file_name} 生成了 {len(chunks)} 个chunks")
            
            # 计算新块的向量
            new_embeddings = self.embedding_model.get_embeddings(chunks)
            
            # 如果索引不存在，创建新索引
            if self.index_manager.index is None:
                print("创建新的FAISS索引")
                self.index_manager.create_index(new_embeddings)
            else:
                # 增量添加到现有索引
                print("增量添加到现有索引")
                self.index_manager.index.add(new_embeddings)
            
            # 更新数据结构
            doc_source_idx = len(self.doc_sources)
            self.doc_sources.append(doc_path)
            
            start_idx = len(self.all_chunks)
            self.all_chunks.extend(chunks)
            
            for i in range(len(chunks)):
                self.chunks_to_document[start_idx + i] = doc_source_idx
            
            self.document_to_chunks[doc_source_idx] = chunks
            
            # 异步保存更新后的索引和映射
            await asyncio.get_event_loop().run_in_executor(None, self.save_index_and_mapping)
            
            # 保存文档块到JSON
            self.document_chunker.save_chunks_to_json(chunks, file_name, doc_path, markdown_content)
            
            print(f"文档 {file_name} 异步处理完成")
            return True
            
        except Exception as e:
            print(f"异步处理文档 {doc_path} 失败: {str(e)}")
            return False
    
    def add_document(self, doc_path: str) -> bool:
        """同步版本的文档添加（保持向后兼容）
        
        Args:
            doc_path: 文档路径
            
        Returns:
            是否成功添加
        """
        # 使用asyncio运行异步版本
        try:
            loop = asyncio.get_event_loop()
            return loop.run_until_complete(self.add_document_async(doc_path))
        except RuntimeError:
            # 如果没有事件循环，创建一个新的
            return asyncio.run(self.add_document_async(doc_path))
    
    def validate_index_consistency(self) -> bool:
        """验证索引和chunks的一致性
        
        Returns:
            是否一致
        """
        if self.index_manager.index is None or not self.all_chunks:
            return False
            
        # 检查索引中的向量数量是否与chunks数量一致
        index_count = self.index_manager.index.ntotal
        chunks_count = len(self.all_chunks)
        
        if index_count != chunks_count:
            print(f"索引不一致: 索引中有 {index_count} 个向量，但有 {chunks_count} 个chunks")
            return False
            
        print(f"索引一致性验证通过: {index_count} 个向量对应 {chunks_count} 个chunks")
        return True


# 创建全局文档处理器实例
document_processor = DocumentProcessor()